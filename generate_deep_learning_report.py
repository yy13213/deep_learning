import csv
import json
from pathlib import Path
from typing import Iterable, List, Sequence

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageOps, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
REPORT_DIR = ROOT / "report"
ASSET_DIR = REPORT_DIR / "assets"
REPORT_PATH = REPORT_DIR / "深度学习实验报告.docx"


def ensure_dirs() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: str):
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def make_exp1_comparison_chart() -> Path:
    baseline = load_json("experiment_outputs/cnn_flower_py/summary.json")
    optimized = load_json("experiment_outputs/cnn_flower_optimized/summary.json")
    residual = load_json("experiment_outputs/cnn_flower_residual/summary.json")

    labels = ["Baseline CNN", "Optimized CNN", "Residual CNN"]
    vals = [
        baseline["final_val_acc"],
        optimized["best_val_acc"],
        residual["best_val_acc"],
    ]

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, vals, color=["#9fb3c8", "#5b8ff9", "#61d9a3"])
    ax.set_ylim(0, 1.0)
    ax.set_ylabel("Validation Accuracy")
    ax.set_title("Experiment 1 Accuracy Comparison")
    ax.grid(axis="y", alpha=0.25)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, v + 0.02, f"{v:.3f}", ha="center", va="bottom", fontsize=10)
    fig.tight_layout()
    out = ASSET_DIR / "exp1_accuracy_comparison.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return out


def make_exp2_comparison_chart() -> Path:
    baseline = load_json("experiment_outputs/rnn_sin_py/summary.json")
    optimized = load_json("experiment_outputs/rnn_sin_optimized/summary_optimized.json")

    models = ["RNN", "LSTM", "GRU"]
    base_vals = [
        baseline["rnn_test_mse"],
        baseline["lstm_test_mse"],
        baseline["gru_test_mse"],
    ]
    opt_vals = [optimized["models"][m.lower()]["metrics_norm"]["mse"] for m in models]

    x = np.arange(len(models))
    width = 0.35
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(x - width / 2, base_vals, width, label="Baseline", color="#7cb5ec")
    ax.bar(x + width / 2, opt_vals, width, label="Optimized", color="#90ed7d")
    ax.set_xticks(x)
    ax.set_xticklabels(models)
    ax.set_yscale("log")
    ax.set_ylabel("Test MSE (log scale)")
    ax.set_title("Experiment 2 Test MSE Comparison")
    ax.legend()
    ax.grid(axis="y", alpha=0.25, which="both")
    fig.tight_layout()
    out = ASSET_DIR / "exp2_mse_comparison.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return out


def make_exp3_comparison_chart() -> Path:
    baseline = load_json("experiment_outputs/dcgan_mnist_py/summary.json")
    improved = load_json("experiment_outputs/dcgan_mnist_improved/summary.json")
    labels = ["D loss", "G loss"]
    base_vals = [baseline["final_D_loss"], baseline["final_G_loss"]]
    imp_vals = [improved["final_D_loss"], improved["final_G_loss"]]

    x = np.arange(len(labels))
    width = 0.35
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(x - width / 2, base_vals, width, label="Baseline DCGAN", color="#f6bd16")
    ax.bar(x + width / 2, imp_vals, width, label="Improved DCGAN", color="#5ad8a6")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Final loss")
    ax.set_title("Experiment 3 Final GAN Loss Comparison")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out = ASSET_DIR / "exp3_loss_comparison.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return out


def make_snapshot_montage(image_paths: Sequence[Path], titles: Sequence[str], out_path: Path, cell_w: int = 260) -> Path:
    images = []
    for p in image_paths:
        img = Image.open(p).convert("RGB")
        ratio = cell_w / img.width
        img = img.resize((cell_w, int(img.height * ratio)))
        images.append(img)
    heights = [img.height for img in images]
    title_h = 35
    canvas = Image.new("RGB", (cell_w * len(images), max(heights) + title_h), "white")
    draw = ImageDraw.Draw(canvas)
    for i, (img, title) in enumerate(zip(images, titles)):
        x = i * cell_w
        canvas.paste(img, (x, title_h))
        draw.text((x + 8, 8), title, fill="black")
    canvas.save(out_path)
    return out_path


def csv_top_bottom_table(csv_path: Path, top_n: int = 5):
    with csv_path.open("r", encoding="utf-8") as f:
        reader = csv.reader(f)
        header = next(reader)[1:]
        rows = list(reader)
    diag = []
    for row in rows:
        label = row[0]
        vals = list(map(int, row[1:]))
        idx = header.index(label)
        total = sum(vals)
        acc = vals[idx] / total if total else 0.0
        diag.append((label, acc, total))
    diag_sorted = sorted(diag, key=lambda x: x[1], reverse=True)
    return diag_sorted[:top_n], diag_sorted[-top_n:]


def configure_doc_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_picture(doc: Document, path: Path, width: float = 5.8, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    if caption:
        add_caption(doc, caption)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_report() -> Path:
    ensure_dirs()

    exp1_base = load_json("experiment_outputs/cnn_flower_py/summary.json")
    exp1_opt = load_json("experiment_outputs/cnn_flower_optimized/summary.json")
    exp1_res = load_json("experiment_outputs/cnn_flower_residual/summary.json")
    exp2_base = load_json("experiment_outputs/rnn_sin_py/summary.json")
    exp2_opt = load_json("experiment_outputs/rnn_sin_optimized/summary_optimized.json")
    exp3_base = load_json("experiment_outputs/dcgan_mnist_py/summary.json")
    exp3_imp = load_json("experiment_outputs/dcgan_mnist_improved/summary.json")

    exp1_compare = make_exp1_comparison_chart()
    exp2_compare = make_exp2_comparison_chart()
    exp3_compare = make_exp3_comparison_chart()

    base_montage = make_snapshot_montage(
        [
            ROOT / "experiment_outputs/dcgan_mnist_py/snapshot_epoch_001.png",
            ROOT / "experiment_outputs/dcgan_mnist_py/snapshot_epoch_005.png",
            ROOT / "experiment_outputs/dcgan_mnist_py/snapshot_epoch_010.png",
        ],
        ["Baseline epoch 1", "Baseline epoch 5", "Baseline epoch 10"],
        ASSET_DIR / "dcgan_baseline_montage.png",
    )
    imp_montage = make_snapshot_montage(
        [
            ROOT / "experiment_outputs/dcgan_mnist_improved/snapshot_epoch_001.png",
            ROOT / "experiment_outputs/dcgan_mnist_improved/snapshot_epoch_010.png",
            ROOT / "experiment_outputs/dcgan_mnist_improved/snapshot_epoch_030.png",
            ROOT / "experiment_outputs/dcgan_mnist_improved/snapshot_epoch_050.png",
        ],
        ["Improved 1", "Improved 10", "Improved 30", "Improved 50"],
        ASSET_DIR / "dcgan_improved_montage.png",
        cell_w=210,
    )

    opt_top, opt_bottom = csv_top_bottom_table(ROOT / "experiment_outputs/cnn_flower_optimized/best_confusion_matrix.csv")
    res_top, res_bottom = csv_top_bottom_table(ROOT / "experiment_outputs/cnn_flower_residual/best_confusion_matrix.csv")

    doc = Document()
    configure_doc_style(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    add_title(doc, "深度学习实验报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("内容覆盖实验一 CNN 图像分类、实验二 RNN/LSTM 时间序列预测、实验三 DCGAN/BIGGAN 图像生成\n")
    p.add_run("依据代码（Python 版本）、实验任务书与 experiment_outputs 中已有实验记录自动整理生成")

    doc.add_heading("1. 实验概述", level=1)
    doc.add_paragraph(
        "本报告根据仓库中的 Python 实验脚本、三份 PDF 实验任务书，以及当前目录下保存的实验结果文件撰写。"
        "报告的组织原则为：先按照任务书要求说明基础实验的目的、原理、流程与结果，再介绍我们对深度神经网络结构和训练策略所做的改进，并结合实验记录进行对比分析。"
    )
    doc.add_paragraph(
        "需要说明的是，实验一任务书文本描述使用 Kaggle Flower Recognition 五分类数据集，而当前实验输出文件中的 `class_to_idx` 实际包含 102 个类别；"
        "因此本报告在原理和步骤部分仍按任务书思路说明，在结果分析部分则严格采用实际运行记录。"
    )

    doc.add_heading("2. 实验环境", level=1)
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["操作系统", "Linux 6.8.0 x86_64"],
            ["Python", "3.10.20"],
            ["PyTorch", "2.10.0+cu128"],
            ["torchvision", "0.25.0+cu128"],
            ["GPU", "2 × NVIDIA RTX 6000 Ada Generation (49 GB)"],
            ["驱动版本", "570.211.01"],
        ],
    )

    doc.add_heading("3. 实验一：卷积神经网络图像分类", level=1)
    doc.add_heading("3.1 任务书要求与基础原理", level=2)
    doc.add_paragraph(
        "实验一要求利用 Python 与 PyTorch 搭建卷积神经网络完成花卉图像分类。任务书的核心流程包括："
        "数据预处理（Resize、ToTensor、Normalize）、自定义 Dataset、按比例划分训练集与验证集、"
        "构建卷积层 + 激活层 + 池化层 + 全连接层的 CNN 结构，并用交叉熵损失进行监督训练。"
    )
    add_bullets(
        doc,
        [
            "输入层负责图像标准化与尺寸统一。",
            "卷积层通过局部感受野和权值共享抽取纹理、边缘和语义模式。",
            "激活层引入非线性，使模型具备更强的表达能力。",
            "池化层用于降维、增强平移容忍度并抑制过拟合。",
            "全连接层输出最终类别 logits，并通过 CrossEntropyLoss 完成训练。",
        ],
    )

    doc.add_heading("3.2 基础实现与结果", level=2)
    doc.add_paragraph(
        "基础版脚本 `cnn_flower_classification.py` 采用 28×28 输入、两层卷积和三层全连接的浅层 CNN。"
        f"实验记录显示：训练 {exp1_base['epochs']} 个 epoch，batch size 为 {exp1_base['batch_size']}，学习率为 {exp1_base['learning_rate']}，最终验证准确率为 {exp1_base['final_val_acc']:.4f}。"
    )
    add_picture(doc, ROOT / "experiment_outputs/cnn_flower_py/fig_baseline_loss_acc.png", width=5.9, caption="图 1 实验一基础版 CNN 的损失与准确率曲线")

    doc.add_heading("3.3 模型改进：优化版 CNN", level=2)
    doc.add_paragraph(
        "在 `cnn_flower_classification_optimized.py` 中，我们对基础版进行了较大幅度优化。主要改动包括："
        "将输入分辨率提升到 224×224、使用 ImageNet 归一化统计量、引入训练/验证不同的数据变换、"
        "加入随机翻转、旋转、颜色扰动与随机擦除等增强策略，并将浅层 CNN 替换为带 BatchNorm、Dropout 和更宽通道数的深层 CNN。"
    )
    add_bullets(
        doc,
        [
            "输入尺寸由 28×28 提升到 224×224，更适合自然图像细节保留。",
            "优化器改为 AdamW，并加入权重衰减、梯度裁剪与余弦退火学习率调度。",
            "损失函数中加入 label smoothing，缓解模型过度自信。",
            "训练集按类别随机打乱后再做分层切分，降低数据顺序偏差。",
        ],
    )
    doc.add_paragraph(
        f"优化版最佳验证准确率提升到 {exp1_opt['best_val_acc']:.4f}，明显优于基础版的 {exp1_base['final_val_acc']:.4f}。"
        f"其最佳模型出现在第 {exp1_opt['best_epoch']} 个 epoch。"
    )
    add_picture(doc, ROOT / "experiment_outputs/cnn_flower_optimized/fig_optimized_loss_acc.png", width=5.9, caption="图 2 实验一优化版 CNN 的损失与准确率曲线")
    add_picture(doc, ROOT / "experiment_outputs/cnn_flower_optimized/fig_best_confusion_matrix.png", width=5.9, caption="图 3 实验一优化版 CNN 的最佳混淆矩阵")

    doc.add_heading("3.4 模型改进：残差 CNN", level=2)
    doc.add_paragraph(
        "进一步地，`cnn_flower_classification_residual.py` 在优化版 CNN 的基础上引入 ResidualBlock，"
        "通过 shortcut 结构缓解深层网络训练困难，使网络学习残差映射而非完整映射。"
        "残差结构特别适合更深的卷积主干，能更稳定地抽取中高层纹理特征。"
    )
    doc.add_paragraph(
        f"残差版最佳验证准确率达到 {exp1_res['best_val_acc']:.4f}，在现有记录中优于优化版 CNN 的 {exp1_opt['best_val_acc']:.4f}，"
        f"最佳轮次为第 {exp1_res['best_epoch']} 个 epoch。"
    )
    add_picture(doc, ROOT / "experiment_outputs/cnn_flower_residual/fig_residual_loss_acc.png", width=5.9, caption="图 4 实验一残差 CNN 的损失与准确率曲线")
    add_picture(doc, ROOT / "experiment_outputs/cnn_flower_residual/fig_best_confusion_matrix.png", width=5.9, caption="图 5 实验一残差 CNN 的最佳混淆矩阵")
    add_picture(doc, exp1_compare, width=5.8, caption="图 6 实验一基础版、优化版与残差版验证准确率对比")
    add_table(
        doc,
        ["模型", "关键设置", "验证准确率"],
        [
            ["基础版 CNN", "28×28 输入，浅层两卷积", f"{exp1_base['final_val_acc']:.4f}"],
            ["优化版 CNN", "224×224 + 增强 + AdamW + Cosine", f"{exp1_opt['best_val_acc']:.4f}"],
            ["残差 CNN", "优化版基础上加入 ResidualBlock", f"{exp1_res['best_val_acc']:.4f}"],
        ],
    )

    doc.add_paragraph("从类别层面看，残差版在难分类类别上的提升更明显。下表给出优化版和残差版中若干最优/最差类别的准确率。")
    add_table(
        doc,
        ["模型", "类别准确率 Top-5", "类别准确率 Bottom-5"],
        [
            [
                "优化版 CNN",
                "；".join([f"{k}:{v:.2f}" for k, v, _ in opt_top]),
                "；".join([f"{k}:{v:.2f}" for k, v, _ in opt_bottom]),
            ],
            [
                "残差 CNN",
                "；".join([f"{k}:{v:.2f}" for k, v, _ in res_top]),
                "；".join([f"{k}:{v:.2f}" for k, v, _ in res_bottom]),
            ],
        ],
    )
    doc.add_paragraph(
        "综合来看，实验一表明：在花卉图像分类任务中，仅靠浅层 CNN 难以取得较高精度；"
        "通过更高输入分辨率、数据增强、正则化、优化器改进与残差连接，可以显著提高模型的泛化能力。"
    )

    doc.add_heading("4. 实验二：循环神经网络时间序列预测", level=1)
    doc.add_heading("4.1 任务书要求与基础原理", level=2)
    doc.add_paragraph(
        "实验二要求利用正弦或余弦序列构造时间序列预测任务，采用前 3 个采样点预测下一个采样点，"
        "并比较基础 RNN、LSTM 与 GRU 模型。任务书强调理解循环神经网络的时序记忆机制，以及 LSTM 的遗忘门、输入门和输出门结构。"
    )
    add_bullets(
        doc,
        [
            "RNN 将前一时刻隐藏状态传递到下一时刻，适合处理时序数据。",
            "LSTM 通过记忆单元与门控结构缓解长距离依赖问题。",
            "GRU 在结构上更紧凑，参数更少，适合作为 LSTM 的对照模型。",
        ],
    )

    doc.add_heading("4.2 基础实现与结果", level=2)
    doc.add_paragraph(
        "基础版脚本 `rnn_lstm_sin.py` 以 `sin(πx)` 为序列，采用滑动窗口构造样本。"
        "在修复目标张量维度广播问题后，基础版三种模型都能稳定收敛，并在测试集上得到较低误差。"
    )
    add_table(
        doc,
        ["模型", "最终训练损失", "测试 MSE"],
        [
            ["RNN", f"{exp2_base['rnn_final_train_loss']:.6e}", f"{exp2_base['rnn_test_mse']:.6e}"],
            ["LSTM", f"{exp2_base['lstm_final_train_loss']:.6e}", f"{exp2_base['lstm_test_mse']:.6e}"],
            ["GRU", f"{exp2_base['gru_final_train_loss']:.6e}", f"{exp2_base['gru_test_mse']:.6e}"],
        ],
    )
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_py/fig_train_loss_all.png", width=5.8, caption="图 7 实验二基础版 RNN/LSTM/GRU 训练损失曲线")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_py/fig_pred_rnn.png", width=5.8, caption="图 8 实验二基础版 RNN 预测结果")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_py/fig_pred_lstm.png", width=5.8, caption="图 9 实验二基础版 LSTM 预测结果")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_py/fig_pred_gru.png", width=5.8, caption="图 10 实验二基础版 GRU 预测结果")

    doc.add_heading("4.3 模型改进：长窗口 + 更强预测头 + 训练策略优化", level=2)
    doc.add_paragraph(
        "在 `rnn_lstm_sin_optimized.py` 中，我们对实验二进行了系统优化："
        "将 look_back 从 3 提升为 30，引入训练/验证/测试三划分，归一化参数仅用训练集拟合，"
        "将简单线性预测头替换为 `LayerNorm + Dropout + MLP Head`，并增加残差预测、AdamW、梯度裁剪、"
        "学习率自适应衰减和 early stopping。"
    )
    add_bullets(
        doc,
        [
            "长窗口能让模型观察更长的局部周期信息，提升对正弦曲率的拟合能力。",
            "残差预测采用 `ŷ = x_t + Δ`，更符合平滑函数的局部连续性。",
            "验证集与早停机制使训练过程更加规范，也便于后续超参数搜索。",
        ],
    )
    doc.add_paragraph(
        "优化版实验记录表明，LSTM 在标准化空间下取得了最优测试误差，"
        f"其测试 MSE 为 {exp2_opt['models']['lstm']['metrics_norm']['mse']:.6e}，R² 达到 {exp2_opt['models']['lstm']['metrics_norm']['r2']:.6f}。"
    )
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_optimized/fig_loss_train_val_all.png", width=5.8, caption="图 11 实验二优化版各模型训练/验证损失曲线")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_optimized/fig_pred_rnn.png", width=5.8, caption="图 12 实验二优化版 RNN 预测结果")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_optimized/fig_pred_lstm.png", width=5.8, caption="图 13 实验二优化版 LSTM 预测结果")
    add_picture(doc, ROOT / "experiment_outputs/rnn_sin_optimized/fig_pred_gru.png", width=5.8, caption="图 14 实验二优化版 GRU 预测结果")
    add_picture(doc, exp2_compare, width=5.8, caption="图 15 实验二基础版与优化版测试 MSE 对比（对数坐标）")
    add_table(
        doc,
        ["模型", "优化版测试 MSE", "优化版测试 RMSE", "优化版测试 R²"],
        [
            [
                "RNN",
                f"{exp2_opt['models']['rnn']['metrics_norm']['mse']:.6e}",
                f"{exp2_opt['models']['rnn']['metrics_norm']['rmse']:.6e}",
                f"{exp2_opt['models']['rnn']['metrics_norm']['r2']:.6f}",
            ],
            [
                "LSTM",
                f"{exp2_opt['models']['lstm']['metrics_norm']['mse']:.6e}",
                f"{exp2_opt['models']['lstm']['metrics_norm']['rmse']:.6e}",
                f"{exp2_opt['models']['lstm']['metrics_norm']['r2']:.6f}",
            ],
            [
                "GRU",
                f"{exp2_opt['models']['gru']['metrics_norm']['mse']:.6e}",
                f"{exp2_opt['models']['gru']['metrics_norm']['rmse']:.6e}",
                f"{exp2_opt['models']['gru']['metrics_norm']['r2']:.6f}",
            ],
        ],
    )
    doc.add_paragraph(
        "实验二说明：对于简单时间序列任务，基础 RNN 就已经可以实现拟合；但在更规范的划分、更强的预测头和更稳定的训练策略支持下，"
        "LSTM 与 GRU 的优势会更加清晰，其中 LSTM 通常在当前记录中表现最好。"
    )

    doc.add_heading("5. 实验三：DCGAN 与 BIGGAN 图像生成", level=1)
    doc.add_heading("5.1 任务书要求与基础原理", level=2)
    doc.add_paragraph(
        "实验三要求在 MNIST 上实现 DCGAN 手写数字生成，并进一步分析 BIGGAN 的网络特点，"
        "比较不同生成式对抗网络的建模能力。DCGAN 的核心思想是使用卷积和反卷积替代传统 GAN 中的全连接生成器/判别器，"
        "通过 BatchNorm、ReLU、LeakyReLU 等机制改善训练稳定性。"
    )
    add_bullets(
        doc,
        [
            "生成器负责将随机噪声映射成图像，目标是“骗过”判别器。",
            "判别器负责区分真实图像和生成图像，目标是提高识别准确率。",
            "GAN 训练是一个最小最大博弈过程，损失曲线通常存在震荡，因此需要同时查看可视化样本。",
            "BIGGAN 在更大模型规模、类别条件建模与训练技巧上更强，但对数据量和算力要求更高。",
        ],
    )

    doc.add_heading("5.2 基础版 DCGAN 实现与结果", level=2)
    doc.add_paragraph(
        "基础版脚本 `dcgan_biggan_mnist.py` 实现了标准 MNIST DCGAN："
        "使用反卷积生成器和卷积判别器，以 BCE 对抗损失训练。当前实验记录中，基础版训练 10 个 epoch，"
        f"最终判别器损失为 {exp3_base['final_D_loss']:.4f}，生成器损失为 {exp3_base['final_G_loss']:.4f}。"
    )
    add_picture(doc, ROOT / "experiment_outputs/dcgan_mnist_py/fig_train_loss.png", width=5.8, caption="图 16 实验三基础版 DCGAN 训练损失曲线")
    add_picture(doc, base_montage, width=6.0, caption="图 17 实验三基础版 DCGAN 在不同 epoch 的生成结果演化")
    add_picture(doc, ROOT / "experiment_outputs/dcgan_mnist_py/generated_final.png", width=5.2, caption="图 18 实验三基础版 DCGAN 的最终生成结果")

    doc.add_heading("5.3 模型改进：更强判别器、Hinge Loss、EMA 与上采样生成器", level=2)
    doc.add_paragraph(
        "在 `dcgan_mnist_improved.py` 中，我们围绕超参数、网络结构和损失函数三方面进行了增强。"
        "核心思路包括：提升训练轮数到 50，噪声维度增大到 128，生成器采用 `Upsample + Conv` 结构以减少棋盘伪影，"
        "判别器升级为更深的 strong 版本，并加入 SpectralNorm 与 Dropout；"
        "损失函数从 BCE 切换为更稳定的 Hinge Loss，同时使用 EMA 保存更平滑的生成器权重。"
    )
    add_bullets(
        doc,
        [
            "Hinge GAN 与 SpectralNorm 组合通常比 BCE 更稳定。",
            "EMA 生成器在采样时能得到更平滑、更稳定的数字轮廓。",
            "strong 判别器可以提升真假样本区分能力，但需要通过 Dropout 和 instance noise 避免过强。",
        ],
    )
    doc.add_paragraph(
        f"改进版训练 {exp3_imp['epochs']} 个 epoch，最终生成器损失下降到 {exp3_imp['final_G_loss']:.4f}，"
        f"最终判别器损失为 {exp3_imp['final_D_loss']:.4f}。从生成样本上看，改进版数字边缘更平滑、类别模式更多样。"
    )
    add_picture(doc, ROOT / "experiment_outputs/dcgan_mnist_improved/fig_train_loss.png", width=5.8, caption="图 19 实验三改进版 DCGAN 训练损失曲线")
    add_picture(doc, ROOT / "experiment_outputs/dcgan_mnist_improved/fig_discriminator_logits.png", width=5.8, caption="图 20 实验三改进版判别器 logits 变化曲线")
    add_picture(doc, imp_montage, width=6.0, caption="图 21 实验三改进版 DCGAN 在不同 epoch 的生成结果演化")
    add_picture(doc, ROOT / "experiment_outputs/dcgan_mnist_improved/generated_final.png", width=5.2, caption="图 22 实验三改进版 DCGAN 的最终生成结果")
    add_picture(doc, exp3_compare, width=5.8, caption="图 23 实验三基础版与改进版最终损失对比")
    add_table(
        doc,
        ["方案", "关键设置", "最终 D loss", "最终 G loss"],
        [
            ["基础版 DCGAN", "BCE + 基础卷积判别器 + deconv", f"{exp3_base['final_D_loss']:.4f}", f"{exp3_base['final_G_loss']:.4f}"],
            ["改进版 DCGAN", "Hinge + strong D + upsample G + EMA", f"{exp3_imp['final_D_loss']:.4f}", f"{exp3_imp['final_G_loss']:.4f}"],
        ],
    )

    doc.add_heading("5.4 BIGGAN 分析与对比", level=2)
    doc.add_paragraph(
        "任务书要求进一步分析 BIGGAN。当前代码仓库中提供了 BIGGAN 预训练接口和实验分析说明，但 `experiment_outputs` 下并没有独立的 BIGGAN 自训练结果记录；"
        "因此本报告对 BIGGAN 部分以原理分析和与 DCGAN 的对比总结为主，不虚构额外实验结果。"
    )
    add_table(
        doc,
        ["维度", "DCGAN", "BIGGAN"],
        [
            ["训练数据规模", "适合 MNIST 等小规模数据集", "通常面向大规模类别数据集（如 ImageNet）"],
            ["生成方式", "无条件或弱条件生成较常见", "强条件生成能力更突出"],
            ["模型规模", "较小，易训练", "更大、更深，对显存和训练技巧要求高"],
            ["图像质量", "适合教学与入门实验", "在复杂自然图像上更逼真"],
            ["当前仓库证据", "有完整训练记录和生成图", "有代码接口和原理说明，无本地训练输出记录"],
        ],
    )

    doc.add_heading("6. 综合结论", level=1)
    doc.add_paragraph(
        "三个实验的共同结论是：在保证基础模型可运行的前提下，合理的数据预处理、网络结构增强和训练策略优化，能够显著改善深度学习实验效果。"
        "实验一中，优化版 CNN 与残差 CNN 分别将验证准确率从 0.3971 提升到 0.7087 和 0.7735；"
        "实验二中，时间序列任务在规范的数据划分和更强的序列建模头后取得了更高精度，其中优化版 LSTM 的测试 MSE 最优；"
        "实验三中，改进版 DCGAN 在 50 个 epoch 的充分训练后，生成样本的可辨识度与多样性都优于基础版。"
    )
    doc.add_paragraph(
        "进一步看，这些改进分别体现了三类深度学习优化思想："
        "图像分类任务中更深的空间特征提取与残差连接，序列预测任务中更长历史窗口与更稳定的优化器，"
        "以及生成模型任务中对抗损失设计、判别器归一化与 EMA 采样策略。"
        "这些经验对于后续课程实验和更复杂任务迁移都具有直接参考价值。"
    )

    doc.add_heading("附录：对应代码与实验材料", level=1)
    add_bullets(
        doc,
        [
            "实验一基础版：exp/exp1/cnn_flower_classification.py",
            "实验一优化版：exp/exp1/cnn_flower_classification_optimized.py",
            "实验一残差版：exp/exp1/cnn_flower_classification_residual.py",
            "实验二基础版：exp/exp2/rnn_lstm_sin.py",
            "实验二优化版：exp/exp2/rnn_lstm_sin_optimized.py",
            "实验三基础版：exp/exp3/dcgan_biggan_mnist.py",
            "实验三改进版：exp/exp3/dcgan_mnist_improved.py",
            "基础与改进实验图表、摘要文件均来自 experiment_outputs/ 目录。",
        ],
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
